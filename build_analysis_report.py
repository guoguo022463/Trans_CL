# -*- coding: utf-8 -*-
"""生成 v1 / v2 对比分析报告 docx（standard_business_brief 风格）。"""
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, 'Trans-CL_v1_v2_对比分析报告.docx')

FONT_EN = 'Calibri'
FONT_CN = 'Microsoft YaHei'

HEADER_FILL = 'F2F4F7'
INK = '0B2545'
H1_COLOR = '2E74B5'
H2_COLOR = '2E74B5'
H3_COLOR = '1F4D78'


def set_run_font(run, size=None, bold=None, color=None, italic=None, name=FONT_EN):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.font.italic = italic


def para(doc, text='', size=11, after=6, line=1.10, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    tok = {1: (16, H1_COLOR, 16, 8), 2: (13, H2_COLOR, 12, 6), 3: (12, H3_COLOR, 8, 4)}[level]
    size, color, before, after = tok
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=color)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    pf = p.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.167
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def set_cell_bg(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def set_table_geometry(table, col_widths_in):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '9360')
    tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120')
    tblInd.set(qn('w:type'), 'dxa')
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is None:
        tblCellMar = OxmlElement('w:tblCellMar')
        tblPr.append(tblCellMar)
    for tag, val in [('top', 80), ('bottom', 80), ('start', 120), ('end', 120)]:
        el = tblCellMar.find(qn('w:' + tag))
        if el is None:
            el = OxmlElement('w:' + tag)
            tblCellMar.append(el)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    total = sum(int(round(w * 1440)) for w in col_widths_in)
    widths = [int(round(w * 1440 * 9360 / total)) for w in col_widths_in]
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[i]))
            tcW.set(qn('w:type'), 'dxa')


def fill_table(table, rows, header=True, font_size=10.5):
    for ri, row_vals in enumerate(rows):
        row = table.rows[ri]
        for ci, val in enumerate(row_vals):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            pf.line_spacing = 1.0
            if header and ri == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_run_font(r, size=font_size, bold=(header and ri == 0))
            if header and ri == 0:
                set_cell_bg(cell, HEADER_FILL)


def add_table(doc, rows, widths):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    fill_table(t, rows)
    set_table_geometry(t, widths)
    return t


def caption(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=9, color='595959', italic=True)
    return p


def add_header(doc, text):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_run_font(r, size=9, color='7F7F7F')
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), 'C9D4E4')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_footer(doc, label):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run(label + '  |  ')
    set_run_font(r1, size=9, color='7F7F7F')
    r2 = p.add_run('第 ')
    set_run_font(r2, size=9, color='7F7F7F')
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    rr = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18')
    col = OxmlElement('w:color'); col.set(qn('w:val'), '7F7F7F')
    rpr.append(col); rpr.append(sz)
    rr.append(rpr)
    t = OxmlElement('w:t'); t.text = '1'
    rr.append(t)
    fld.append(rr)
    p._p.append(fld)
    r3 = p.add_run(' 页')
    set_run_font(r3, size=9, color='7F7F7F')


def add_rule(doc):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9D4E4')
    pBdr.append(bottom)
    pPr.append(pBdr)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles['Normal']
    normal.font.name = FONT_EN
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), FONT_EN)
    rfonts.set(qn('w:hAnsi'), FONT_EN)
    rfonts.set(qn('w:eastAsia'), FONT_CN)
    npf = normal.paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(6)
    npf.line_spacing = 1.10

    add_header(doc, 'Trans-CL SOC 竞赛模型对比分析')
    add_footer(doc, 'v1 / v2 对比分析报告')

    # ---------- 标题区 ----------
    t = para(doc, 'Trans-CL SOC 竞赛模型 v1 / v2 对比分析报告', size=22, after=4, bold=True, color=INK)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para(doc, '基于论文 Trans-CL 的 SOC 日志三分类实验 · 特征合并与数值稳定修复', size=12, after=2, color=H1_COLOR)
    para(doc, '报告日期：2026-08-26 ｜ 环境：CUDA ｜ 数据：train.parquet（2,056,871 条）',
         size=10, after=4, color='595959')
    add_rule(doc)

    # ---------- 1. 执行摘要 ----------
    heading(doc, '1. 执行摘要', 1)
    para(doc, '在同一口径（内部 85/15 分层验证、相同超参）下，v1（34 维）Macro F1 0.959、'
              'v2（19 维合并）Macro F1 0.948，二者 ACC 均约 0.99，三类 F1 均超过 0.90。'
              '19 维合并以约 1% 的 Macro F1 代价，把模型参数量从 8.87M 降到 4.20M（-53%）。')
    para(doc, '结论：SHAP 驱动的 19 维合并“近似无损”，在保留三类识别能力的同时大幅瘦身；'
              '早前的 v1 失效（Macro F1 0.322）源于 NTXent 数值溢出与外部验证口径差异，'
              '经数值稳定修复后 v1/v2 均已恢复。')

    # ---------- 2. 实验设置 ----------
    heading(doc, '2. 实验设置', 1)
    add_table(doc, [
        ['项目', '说明'],
        ['任务', 'SOC 日志三分类（benign / suspicious / malicious）'],
        ['训练数据', 'train.parquet（2,056,871 条，含 label_binary）'],
        ['v1 特征', '34 维（TF-IDF PCA 16 维 + 时间/IP/关键词/缺失标志 18 维）'],
        ['v2 特征', '19 维（合并版：identity_missing + 时间2 + 分类4 + IP2 + 消息3 + TF-IDF7）'],
        ['v1 验证', '内部 85/15 分层划分（308,531 条，seed=42）'],
        ['v2 验证', '内部 85/15 分层划分（308,531 条，seed=42）'],
        ['论文基线', 'arXiv:2505.08816（IFIP Networking 2025）'],
    ], [2.0, 4.5])
    caption(doc, '表 1  实验设置')

    # ---------- 3. 相对论文的改动 ----------
    heading(doc, '3. 相对论文原文做出的改动', 1)
    para(doc, '论文原版是面向 CICIDS2017 网络流量（pcap）的自监督对比学习入侵检测：输入为“流 = 包序列”，'
              '每个包含 timestamp/size/direction/ip_protocol/tcp_flags 五个字段，seq_len=32；'
              '用 BERT（EmbeddingLayer + 4 层 Transformer + OutputLayer 投影 + CLSLayer 二分类）+ NTXent '
              '先做对比预训练，再用余弦相似度做二分类异常检测。')
    add_table(doc, [
        ['维度', '论文原版', 'v1 / v2 改动'],
        ['输入', 'pcap 包序列（32 tokens × 5 字段）', '表格化 SOC 日志特征向量（34/19 维）'],
        ['任务', '二分类异常检测（benign vs attack）', '三分类监督（benign/suspicious/malicious）'],
        ['分类头', 'CLSLayer：256→1024→256→1 sigmoid', '3 类 softmax + FocalLoss'],
        ['损失', 'NTXent（温度 0.5）', 'NTXent（L2 归一化）+ FocalLoss 类别权重'],
        ['嵌入', '数值列 Linear + 分类列 Embedding(dict 65536)', '每列 Linear(1→16)，分类 id 按数值处理'],
        ['序列建模', '32 个包 token，自注意力有语义', '特征列展平成单 token，自注意力退化为无操作'],
        ['特征工程', 'NFStream 包特征', 'TF-IDF+PCA + 时间/IP/关键词/缺失标志（SHAP 驱动）'],
        ['数值稳定', '无', 'NTXent L2 归一化 + 梯度裁剪 + NaN 检测回滚'],
    ], [1.35, 2.55, 2.6])
    caption(doc, '表 2  相对论文原文的改动对照')
    para(doc, '需要特别说明：v1/v2 复用了论文的 TransformerEncoder + 对比学习 NTXent 框架，'
              '但把整条样本的若干特征列展平成单个 token，因此 Transformer 的自注意力在本实现中'
              '实际退化为“无序列交互”，这是与论文原意最本质的一处差异。', size=10, color='595959')

    # ---------- 4. 指标对比 ----------
    heading(doc, '4. v1 与 v2 指标对比', 1)
    add_table(doc, [
        ['指标', 'v1（34 维，内部验证）', 'v2（19 维，内部验证）'],
        ['ACC', '0.9929', '0.9910'],
        ['Macro F1', '0.9591', '0.9480'],
        ['Weighted F1', '0.9930', '0.9912'],
        ['AUC', '0.9999', '0.9999'],
        ['benign F1', '0.9969', '0.9962'],
        ['suspicious F1', '0.9271', '0.9081'],
        ['malicious F1', '0.9533', '0.9396'],
    ], [2.2, 2.15, 2.15])
    caption(doc, '表 3  v1 与 v2 核心指标对比（同口径：内部 85/15 验证）')

    # ---------- 5. 混淆矩阵 ----------
    heading(doc, '5. 混淆矩阵对比', 1)
    img1 = os.path.join(BASE, 'confusion_matrix_v1.png')
    img2 = os.path.join(BASE, 'confusion_matrix_v2.png')
    if os.path.exists(img1):
        doc.add_picture(img1, width=Inches(5.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption(doc, '图 1  v1 混淆矩阵（34 维特征，内部 85/15 验证）')
    if os.path.exists(img2):
        doc.add_picture(img2, width=Inches(5.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption(doc, '图 2  v2 混淆矩阵（19 维合并特征，内部 85/15 验证）')

    # ---------- 6. 错误分析 ----------
    heading(doc, '6. 剩余错误分析', 1)
    para(doc, '两者均做到“恶意类向 benign 的漏报为 0”，剩余错误集中在可疑与恶意两类之间。'
              'v1（34 维）内部验证误判 2,201 条（错误率 0.71%），v2（19 维）误判 2,776 条（错误率 0.90%）；'
              '19 维合并带来的额外误判主要来自 suspicious 召回（401→512）和 benign→malicious 误报（1,192→1,501）。')
    para(doc, '这说明三分类当前的主要难点是 suspicious 与 malicious 的边界，而非 benign 与威胁的边界。'
              '后续可在损失函数上对 suspicious/malicious 增加更强的类别权重或排序约束。')

    # ---------- 7. 特征合并 ----------
    heading(doc, '7. v2 特征合并（SHAP 驱动）', 1)
    para(doc, '基于 SHAP（TreeExplainer）+ 特征相关性分析，把 34 维压缩到 19 维：')
    bullet(doc, '合并两个缺失标志 dst_host_missing / username_missing 为 identity_missing（二者相关系数 0.9999，且与恶意类强相关）。')
    bullet(doc, '时间特征从 5 维降到 2 维（保留 sin/cos，删 hour_norm/weekend/night）。')
    bullet(doc, 'IP 特征删掉恒为 0 的 is_loopback。')
    bullet(doc, '消息特征删掉 msg_is_missing（其完全由 pipeline 决定）。')
    bullet(doc, 'TF-IDF PCA 从 16 维降到 7 维（删掉近零的 pca_7..15）。')

    # ---------- 8. 结论与建议 ----------
    heading(doc, '8. 结论与建议', 1)
    para(doc, '结论：在同口径下，19 维合并相对 34 维 v1 仅损失约 0.011 Macro F1（0.959→0.948），'
              '却把参数量从 8.87M 降到 4.20M，属于“近似无损”的轻量化。')
    bullet(doc, '精度优先：继续用 v1（34 维）；部署体积/推理成本优先：用 v2（19 维）。')
    bullet(doc, '下一步优先用外部私榜复现，验证两类模型在分布漂移下的泛化。')
    bullet(doc, '监督轮数从 5 提到 10–15：当前 epoch 3–5 已接近平台期，但 suspicious/malicious 边界仍可再优化。')

    # ---------- 附录 ----------
    heading(doc, '9. 附录：产物清单', 1)
    add_table(doc, [
        ['文件', '说明'],
        ['confusion_matrix_v1.png / v2.png', '混淆矩阵图'],
        ['models/transcl_v1/run_report.json', 'v1 实验指标与混淆矩阵'],
        ['models/transcl_v2/v2_5ep/run_report.json', 'v2 实验指标与混淆矩阵'],
        ['soc_feature_encoder_v2.py', '19 维合并版特征编码器'],
        ['train_transcl_v2.py / predict_transcl_v2.py', 'v2 训练 / 推理脚本'],
    ], [3.3, 3.2])

    doc.save(OUT)
    print('saved:', OUT)


if __name__ == '__main__':
    main()

# -*- coding: utf-8 -*-
"""生成 0825_04 全量训练（模式2：内部85/15）实验报告 docx"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------- token map (standard_business_brief) ----------------
PAGE_W = Inches(8.5)
PAGE_H = Inches(11)
MARGIN = Inches(1)
CONTENT_W_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN = {'top': 80, 'bottom': 80, 'start': 120, 'end': 120}
HEADER_FILL = 'F2F4F7'
H1 = dict(size=16, color='2E74B5', before=16, after=8)
H2 = dict(size=13, color='2E74B5', before=12, after=6)
H3 = dict(size=12, color='1F4D78', before=8, after=4)
BODY = dict(size=11, after=6, line=1.10)
FONT_EN = 'Calibri'
FONT_CN = 'Microsoft YaHei'

BASE = os.path.dirname(os.path.abspath(__file__))
SAVE_DIR = os.path.join(BASE, 'models', 'transcl_v1', '0825_04')
OUT = os.path.join(BASE, '8.25_0825_04实验报告.docx')


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_EN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.font.italic = italic


def para(doc, text='', size=BODY['size'], after=BODY['after'], line=BODY['line'],
         bold=False, color=None, align=None, style=None):
    p = doc.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    tok = {1: H1, 2: H2, 3: H3}[level]
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(tok['before'])
    pf.space_after = Pt(tok['after'])
    pf.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, size=tok['size'], bold=True, color=tok['color'])
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    pf = p.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.167
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_run_font(r, size=BODY['size'])
    return p


def set_cell_bg(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def set_table_geometry(table, col_widths_in):
    """固定 DXA 表格几何：tblW / tblInd / tblGrid / tcW 一致"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    # width
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(CONTENT_W_DXA))
    tblW.set(qn('w:type'), 'dxa')
    # indent
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(qn('w:w'), str(TABLE_INDENT_DXA))
    tblInd.set(qn('w:type'), 'dxa')
    # cell margins
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is None:
        tblCellMar = OxmlElement('w:tblCellMar')
        tblPr.append(tblCellMar)
    for tag, val in [('top', 80), ('bottom', 80), ('start', 120), ('end', 120)]:
        el = tblCellMar.find(qn('w:' + tag))
        if el is None:
            el = OxmlElement('w:' + tag)
            tblCellMar.append(el)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
    # grid
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    total = 0
    for w_in in col_widths_in:
        total += int(round(w_in * 1440))
    # 将宽度按比例缩放到 CONTENT_W_DXA
    widths = [int(round(w_in * 1440 * CONTENT_W_DXA / total)) for w_in in col_widths_in]
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)
    # per-cell width
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(widths[i]))
            tcW.set(qn('w:type'), 'dxa')


def fill_table(table, rows, header=True, col_align=None, header_fill=HEADER_FILL):
    """rows: list of list[str]; col_align: list of align for data cells"""
    for ri, row_vals in enumerate(rows):
        row = table.rows[ri]
        for ci, val in enumerate(row_vals):
            cell = row.cells[ci]
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            pf.line_spacing = 1.0
            if header and ri == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif col_align and ci < len(col_align) and col_align[ci] is not None:
                p.alignment = col_align[ci]
            r = p.add_run(str(val))
            set_run_font(r, size=10.5, bold=(header and ri == 0))
            if header and ri == 0:
                set_cell_bg(cell, header_fill)


def add_page_number_footer(doc, label):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run(label + '  |  ')
    set_run_font(r1, size=9, color='7F7F7F')
    r2 = p.add_run('第 ')
    set_run_font(r2, size=9, color='7F7F7F')
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    rr = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18')
    col = OxmlElement('w:color'); col.set(qn('w:val'), '7F7F7F')
    rpr.append(col); rpr.append(sz)
    rr.append(rpr)
    t = OxmlElement('w:t'); t.text = '1'
    rr.append(t)
    fld.append(rr)
    p._p.append(fld)
    r3 = p.add_run(' 页')
    set_run_font(r3, size=9, color='7F7F7F')


def add_header(doc, text):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    set_run_font(r, size=9, color='7F7F7F')
    # 下边框
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), 'C9D4E4')
    pBdr.append(bottom)
    pPr.append(pBdr)


def caption(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=9, color='595959', italic=True)
    return p


def main():
    doc = Document()
    # 页面设置
    section = doc.sections[0]
    section.page_width = PAGE_W
    section.page_height = PAGE_H
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Normal 样式
    normal = doc.styles['Normal']
    normal.font.name = FONT_EN
    normal.font.size = Pt(BODY['size'])
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), FONT_EN)
    rfonts.set(qn('w:hAnsi'), FONT_EN)
    rfonts.set(qn('w:eastAsia'), FONT_CN)
    npf = normal.paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(6)
    npf.line_spacing = 1.10

    add_header(doc, 'Trans-CL SOC 竞赛实验报告 · 0825_04')
    add_page_number_footer(doc, '0825_04 全量训练')

    # ================= 标题区 =================
    t = para(doc, '0825_04 全量训练实验报告', size=22, after=4, bold=True, color='0B2545')
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para(doc, '模式2：训练集内部 85/15 分层划分 · Trans-CL 对比学习 + 监督分类', size=12, after=2, color='2E74B5')
    para(doc, '实验日期：2026-08-25 ｜ 运行环境：CUDA ｜ 实验编号：0825_04（自动命名）',
         size=10, after=2, color='595959')
    para(doc, 'W&B 运行：https://wandb.ai/2743739349-/transcl-soc/runs/9oy6uc25',
         size=10, after=8, color='595959')

    # ================= 1. 实验概述 =================
    heading(doc, '1. 实验概述', 1)
    para(doc, '本次实验基于 Trans-CL（Transformer + Contrastive Learning）框架，在第二届浙江省大学生人工智能竞赛'
              '「基于 SOC 日志的网络安全威胁检测算法设计与实现」赛题数据上，验证内部 85/15 分层划分模式的'
              '全量训练流程。训练集为官方 train.parquet（2,056,871 条有标签日志），按 85%/15% 分层切分为训练与'
              '验证子集，用于在训练集内部完成模型调参与效果评估。')
    para(doc, '训练分为两个阶段：阶段一为对比学习（NTXent），以数据增强构造正样本对，学习日志语义表示；'
              '阶段二为三分类监督训练（Focal Loss 缓解类别不平衡），在冻结骨干网络的基础上微调分类头与投影层。')
    para(doc, '实验最终未达到预期效果：对比学习在第 8 轮出现 Loss=NaN 数值溢出，监督阶段全程 Loss=NaN，'
              '模型退化为「全量预测为 benign」，验证集 Macro F1 仅 0.3201。本报告记录实验配置、过程、结果，'
              '并给出根因分析与修复建议。')

    # ================= 2. 数据与划分 =================
    heading(doc, '2. 数据与划分', 1)
    para(doc, '使用 train.parquet 全量数据（2,056,871 条），按类别分层随机切分，验证比例 15%，随机种子 42。'
              '训练集与验证集的类别分布保持一致，符合原始数据 92.4% / 2.2% / 5.4% 的比例结构。')
    caption(doc, '表 1  数据划分结果（0825_04）')
    tbl1 = doc.add_table(rows=4, cols=4)
    tbl1.alignment = WD_TABLE_ALIGNMENT.LEFT
    fill_table(tbl1, [
        ['数据集', 'benign', 'suspicious', 'malicious'],
        ['Train（85%）', '1,614,764（92.4%）', '38,607（2.2%）', '94,969（5.4%）'],
        ['Valid（15%）', '284,959（92.4%）', '6,813（2.2%）', '16,759（5.4%）'],
        ['合计', '1,899,723', '45,420', '111,728'],
    ])
    set_table_geometry(tbl1, [2.3, 1.4, 1.4, 1.4])
    para(doc, '', after=2)

    # ================= 3. 模型与训练配置 =================
    heading(doc, '3. 模型与训练配置', 1)
    para(doc, '模型沿用论文 Trans-CL 结构：32 维 SOC 特征编码（TF-IDF + PCA + 时间/IP/关键词特征）、'
              'Transformer Encoder（4 层、4 头、隐层 256）、128 维投影层与三分类头，参数量 8.87M。')
    caption(doc, '表 2  关键超参数')
    tbl2 = doc.add_table(rows=13, cols=2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    fill_table(tbl2, [
        ['超参数', '取值'],
        ['数据划分', '内部 85/15 分层，seed=42'],
        ['Batch size', '256'],
        ['对比学习轮数 / 学习率', '10 / 1e-4'],
        ['NTXent 温度', '0.5'],
        ['监督轮数 / 学习率', '20 / 1e-4'],
        ['Focal Loss gamma / alpha', '2.0 / [1, 5, 8]'],
        ['早停 patience', '6'],
        ['Transformer 结构', '4 层 · 4 头 · 隐层 256 · dropout 0.1'],
        ['特征维度', '32（TF-IDF(500)→PCA(16) + 16 手工特征）'],
        ['权重衰减', '0'],
        ['优化器', 'Adam'],
        ['模型参数量', '8,872,579（8.87M）'],
    ])
    set_table_geometry(tbl2, [2.7, 3.8])
    para(doc, '', after=2)

    # ================= 4. 训练过程 =================
    heading(doc, '4. 训练过程', 1)
    heading(doc, '4.1 阶段一：对比学习', 2)
    para(doc, '对比学习前 7 轮 Loss 稳定下降（4.4902 → 4.3909），第 8 轮起 Loss 突变为 NaN，'
              '说明模型在该轮发生数值溢出（inf/NaN），后续轮次权重已被污染。')
    img1 = os.path.join(SAVE_DIR, 'loss_curves_0825_04.png')
    if os.path.exists(img1):
        doc.add_picture(img1, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption(doc, '图 1  两阶段 Loss / Acc 曲线（对比学习 Ep8 起 NaN；监督阶段退化）')
    heading(doc, '4.2 阶段二：监督分类', 2)
    para(doc, '监督阶段 7 个 epoch 的训练 Loss 全部为 NaN，train_acc 恒为 0.9236；验证集所有样本被预测为 benign，'
              'confusion matrix 中 suspicious 与 malicious 两列全为 0。第 7 轮触发早停（6 轮无提升），'
              '最终保存的 best checkpoint 为 Epoch 1（Weighted F1=0.8869）。')
    para(doc, '监督阶段 Acc 恒等于验证集 benign 占比（92.4%），是典型的「全预测多数类」退化行为：'
              '对比学习阶段产生的 NaN 权重使 logits 全为 NaN，argmax 统一落到索引 0（benign）。')

    # ================= 5. 评估结果 =================
    heading(doc, '5. 评估结果', 1)
    para(doc, '以下为 best checkpoint（Epoch 1）在内部验证集（308,531 条）上的评估结果。'
              '由于模型实际输出全部为 benign，宏观指标接近「只预测多数类」的基线。')
    caption(doc, '表 3  混淆矩阵（Pred_B / Pred_S / Pred_M）')
    tbl3 = doc.add_table(rows=4, cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT
    fill_table(tbl3, [
        ['真实 \\ 预测', 'Pred_B', 'Pred_S', 'Pred_M'],
        ['True_Benign', '284,959', '0', '0'],
        ['True_Suspicious', '6,813', '0', '0'],
        ['True_Malicious', '16,759', '0', '0'],
    ])
    set_table_geometry(tbl3, [2.0, 1.5, 1.5, 1.5])
    para(doc, '', after=2)

    caption(doc, '表 4  二级指标：TPR / FPR / TNR / FNR')
    tbl4 = doc.add_table(rows=5, cols=5)
    tbl4.alignment = WD_TABLE_ALIGNMENT.LEFT
    fill_table(tbl4, [
        ['类别', 'TPR', 'FPR', 'TNR', 'FNR'],
        ['Benign', '1.0000', '1.0000', '0.0000', '0.0000'],
        ['Suspicious', '0.0000', '0.0000', '1.0000', '1.0000'],
        ['Malicious', '0.0000', '0.0000', '1.0000', '1.0000'],
        ['Macro', '0.3333', '0.3333', '0.6667', '0.6667'],
    ])
    set_table_geometry(tbl4, [1.5, 1.25, 1.25, 1.25, 1.25])
    para(doc, '', after=2)

    caption(doc, '表 5  三级指标：PRE / SEN / F1')
    tbl5 = doc.add_table(rows=5, cols=4)
    tbl5.alignment = WD_TABLE_ALIGNMENT.LEFT
    fill_table(tbl5, [
        ['类别', 'PRE', 'SEN', 'F1'],
        ['Benign', '0.9236', '1.0000', '0.9603'],
        ['Suspicious', '0.0000', '0.0000', '0.0000'],
        ['Malicious', '0.0000', '0.0000', '0.0000'],
        ['Macro', '0.3079', '0.3333', '0.3201'],
    ])
    set_table_geometry(tbl5, [1.8, 1.5, 1.5, 1.5])
    para(doc, '', after=2)

    para(doc, '总体：ACC = 0.9236，Macro F1 = 0.3201，Weighted F1 = 0.8869，AUC = N/A。'
              'Weighted F1 因多数类占比高而虚高，不代表模型对少数威胁具备识别能力；'
              'Macro F1 与 suspicious/malicious 两类指标全部为 0，说明模型完全失效。')
    img2 = os.path.join(SAVE_DIR, 'cm_best.png')
    if os.path.exists(img2):
        doc.add_picture(img2, width=Inches(4.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption(doc, '图 2  验证集混淆矩阵（Epoch 1，全预测 benign）')

    # ================= 6. 问题分析 =================
    heading(doc, '6. 问题分析：Loss 变 NaN 的根因链', 1)
    para(doc, '本次实验失败的直接表现是 Loss=NaN 与模型退化为全 benign 预测，其根因链条如下：')
    bullet(doc, '对比学习数值溢出：NTXent 中 sim = z·zᵀ / temperature（温度 0.5），投影向量 z 未做 L2 归一化。'
                '当某 batch 特征范数增大时，相似度矩阵出现极大值甚至 inf，exp/softmax 计算溢出，Loss 变为 NaN。'
                '3 万条 smoke 测试未触发（数据规模小、范数未达临界值），全量 175 万样本在第 8 轮触发。')
    bullet(doc, '无梯度裁剪：训练循环未对梯度做 clip_grad_norm_，Adam 在 512×512 相似度矩阵上的梯度量级较大，'
                '单 batch 的极端输入即可将权重推向 NaN。')
    bullet(doc, 'NaN 权重跨阶段污染：对比学习权重变 NaN 后直接进入监督阶段，logits 全为 NaN，Focal Loss 也为 NaN；'
                'argmax(NaN) 在 CUDA 上稳定返回索引 0，于是所有样本都被判为 benign。')
    bullet(doc, '早停与选优指标失效：Weighted F1 被 92.4% 的 benign 主导，无法反映少数类学习情况；'
                'best checkpoint 停留在 Epoch 1——那是权重尚未被 NaN 污染的唯一正常快照，'
                '这正是「Loss 在下降但 Epoch 1 最优」矛盾的根本解释。')
    para(doc, '补充：监督阶段第 1 轮 Acc 即 0.9236 且全程不变，进一步证明问题并非监督训练本身，'
              '而是上游对比学习阶段产生的 NaN 权重所致。')

    # ================= 7. 改进建议 =================
    heading(doc, '7. 改进建议', 1)
    bullet(doc, '数值稳定性（优先）：NTXent 前对 z1/z2 执行 F.normalize(p, dim=1)，使相似度限定在 [-1, 1] 内，'
                '从根源避免 exp 溢出；或使用 logsumexp 稳定形式计算 InfoNCE 损失。')
    bullet(doc, '梯度裁剪：对比学习与监督阶段均增加 torch.nn.utils.clip_grad_norm_(params, max_norm=1.0)。')
    bullet(doc, 'NaN 防护：每轮检查 loss 与模型权重是否有限（torch.isfinite），一旦检测到 NaN 立即告警并'
                '回滚该轮更新（或跳过该 batch），绝不将 NaN 权重写入 checkpoint。')
    bullet(doc, '超参数调整：可尝试降低对比学习学习率（如 3e-5）、提高温度（如 1.0），或减少投影维度，'
                '降低梯度量级。')
    bullet(doc, '选优指标：早停与 best 选择改用 Macro F1（或 minority F1），避免 Weighted F1 掩盖退化；'
                '同时监控 suspicious/malicious 两类的 Recall。')
    bullet(doc, 'Checkpoint 校验：保存 best.pth / final.pth 前检查 state_dict 无 NaN；'
                '若检测到污染则回退到最近一次正常权重。')

    # ================= 8. 附录 =================
    heading(doc, '8. 附录', 1)
    heading(doc, '8.1 运行命令', 2)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    r = p.add_run('python -u train_transcl_v1.py --data-path main/data/competition/train.parquet --split 0.15')
    set_run_font(r, size=10, color='404040')
    heading(doc, '8.2 产物文件', 2)
    caption(doc, '表 6  实验产物（models/transcl_v1/0825_04/）')
    tbl6 = doc.add_table(rows=8, cols=2)
    tbl6.alignment = WD_TABLE_ALIGNMENT.LEFT
    fill_table(tbl6, [
        ['文件', '说明'],
        ['final.pth', '最终权重（best epoch 1，best 权重回填）'],
        ['best.pth', '验证集最优 checkpoint'],
        ['contrastive.pth', '对比学习阶段权重'],
        ['encoder.pkl', '特征编码器'],
        ['metrics.csv', '逐 epoch 训练/验证指标'],
        ['run_report.json', '实验配置与最佳指标汇总'],
        ['cm_best.png / loss_curves_0825_04.png', '混淆矩阵图 / 两阶段曲线图'],
    ])
    set_table_geometry(tbl6, [2.6, 3.9])
    para(doc, '', after=2)
    para(doc, '说明：本实验为失败诊断实验，结果为负数（模型退化），不用于正式提交；'
              '修复 NaN 问题后需重新训练验证。', size=10, color='595959')

    doc.save(OUT)
    print('saved:', OUT)


if __name__ == '__main__':
    main()

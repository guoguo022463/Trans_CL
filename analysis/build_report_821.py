import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHART_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "charts_821")
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "8.21_TransCL_SOC实验分析报告.docx")

CN_FONT = "SimSun"          # body Chinese face (宋体)
CN_HEAD = "SimHei"          # headings Chinese face (黑体)

def set_run_font(run, size=11, bold=False, east=CN_FONT, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), east)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if color:
        run.font.color.rgb = color

def add_para(doc, text, size=11, bold=False, align="left", east=CN_FONT,
             space_after=6, space_before=0, line=1.3, indent=False):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if indent:
        pf.first_line_indent = Pt(21)  # 首行缩进 2 字符
    r = p.add_run(text)
    set_run_font(r, size, bold, east)
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12.5}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level == 1 else 10)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, sizes[level], True, CN_HEAD)
    return p

def add_table(doc, rows, header_row=True, widths=None, font_size=9.5):
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 固定布局，避免 Word 自动调整列宽
    tbl.autofit = False
    tblPr = tbl._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(6.5 * 1440)))
    tblPr.append(tblW)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_after = Pt(2)
            pf.space_before = Pt(2)
            r = p.add_run(str(val))
            is_head = (i == 0 and header_row)
            set_run_font(r, font_size, is_head, CN_HEAD)
    if widths:
        for j, wd in enumerate(widths):
            for i in range(len(rows)):
                if j < ncols:
                    tbl.cell(i, j).width = Inches(wd)
    return tbl

def add_figure(doc, path, caption, width=5.6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    rr = cap.add_run(caption)
    set_run_font(rr, 9, False, CN_HEAD)

def shade_header(tbl):
    for cell in tbl.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "D9E1F2")
        tcPr.append(shd)

def bullet(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run("• " + text)
    set_run_font(r, size, False, CN_FONT)
    return p

def num_para(doc, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    set_run_font(r, size, False, CN_FONT)
    return p

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.27)
sec.page_height = Inches(11.69)
for attr in ("left_margin", "right_margin"):
    setattr(sec, attr, Inches(0.8))
sec.top_margin = Inches(0.9)
sec.bottom_margin = Inches(0.9)

# ---------------- 封面色块 / 标题 ----------------
add_para(doc, "Trans-CL SOC", 22, True, "center", CN_HEAD, space_after=6, space_before=6)
add_para(doc, "8.21 实验报告", 17, True, "center", CN_HEAD, space_after=4)
add_para(doc, "自动反频率α + 100%/100%无划分 — 异常运行诊断", 13, False, "center",
         CN_FONT, space_after=4)
add_para(doc, "网络安全防御技术 | 2026-08-22", 11, False, "center", CN_FONT, space_after=14)

# ---------------- 目录 ----------------
add_heading(doc, "目 录", 1)
toc = ["一、实验概述", "二、实验配置", "三、数据集分析", "四、训练过程",
       "4.1 对比学习阶段", "4.2 监督学习阶段", "五、评估结果分析", "5.1 指标概览",
       "5.2 混淆矩阵分析", "六、异常与根因诊断", "6.1 数据划分：100%/100% 数据泄漏",
       "6.2 特征编码：PCA 仅 38.8% 方差", "6.3 对比学习 NaN 发散", "6.4 监督学习 Loss=NaN 与多数类退化",
       "6.5 W&B 步数回绕", "6.6 自动反频率α = 已废弃配置", "七、风险与影响",
       "八、实验结论", "九、总结与建议"]
for t in toc:
    add_para(doc, t, 11, False, "left", CN_FONT, space_after=2, line=1.25)

# ================= 一、实验概述 =================
add_heading(doc, "一、实验概述", 1)
add_para(doc, "本实验为 8.21 运行，目标是使用 Trans-CL SOC 的完整两阶段流程（对比学习 + 监督分类）"
              "在所有 SOC 日志数据上训练三分类模型。与 8.19 的成功运行不同，本次运行出现明显异常："
              "对比学习阶段在第 9-10 epoch 损失变为 NaN，监督学习阶段自第 1 epoch 起训练损失即为 NaN，"
              "模型最终退化为仅预测多数类（Benign）的“哑模型”，Suspicious 与 Malicious 识别完全失效。",
        11, False, "justify", CN_FONT, indent=True)
add_para(doc, "实验配置要点（与 8.19 回退版的主要差异）：", 11, True)
num_para(doc, "(1) 对比学习：10 epochs，NTXent loss，T=0.5；损失在 Ep9-10 变成 NaN，训练数值发散；")
num_para(doc, "(2) 监督学习：20 epochs max，FocalLoss (gamma=2.0)，alpha 为自动反频率加权 "
              "[0.36, 15.09, 6.14]（非 8.19 的硬编码 [1,20,8]），Early Stop patience=6；")
num_para(doc, "(3) 特征编码：32D（Time 5D + Categorical 4D + IP 3D + Message 4D + TF-IDF PCA 16D），"
              "PCA 仅保留 38.8% 方差；")
num_para(doc, "(4) 数据划分：100%/100% 无划分（train 与 valid 均为全部 2,056,871 条），存在数据泄漏，验证集失去意义；")
num_para(doc, "(5) 评估结果：ACC=0.9236（恰等于 Benign 占比 92.4%），Suspicious/Malicious F1=0，AUC=N/A，"
              "表明模型未学到任何有效分类边界。")
add_para(doc, "核心结论：本次运行以失败告终。已保存的“最佳”模型（Ep1）实为退化的多数类分类器，不具备参考价值。"
              "NaN 根因源自对比学习发散，同时配置偏离 8.19 已验证基线，建议回退并修复数值稳定性后重训。",
        11, True, "justify", CN_FONT, indent=True)
add_para(doc, "重点备份（代码 & 模型权重）：本次产物保存于 models/transcl_v4_819_full，与 8.19 目录相近，"
              "存在覆盖已知有效权重的风险，需先行备份。", 10.5, False, "justify", CN_FONT)

# ================= 二、实验配置 =================
add_heading(doc, "二、实验配置", 1)
add_para(doc, "本实验配置在四处关键参数上与 8.19 回退版（正常）不同，其中数据划分与类别权重α是导致失败的主因。"
              "以下为 8.19 回退版（正常）与 8.21 运行（异常）的对比：", 11, False, "justify", CN_FONT, indent=True)
add_para(doc, "表2-1：8.19回退版 / 8.21运行 配置对比", 10.5, True, "center", CN_HEAD)
cfg = [
    ["配置项", "8.19回退版（正常）", "8.21运行（异常/当前）", "差异说明"],
    ["特征维度", "32D (PCA 16)", "32D (PCA 16)", "相同；两者 PCA 保留方差均偏低"],
    ["数据划分", "85/15 分层", "100%/100% 无划分", "8.21 用全量做 train+valid，数据泄漏"],
    ["类别权重α", "硬编码 [1,20,8]", "自动反频率 [0.36,15.09,6.14]", "8.21 自动计算，等价 8.19 已废弃优化版"],
    ["对比学习Epoch", "10", "10", "相同"],
    ["监督学习Epoch", "20 (Early Stop p=6)", "20 (Early Stop p=6)", "相同"],
    ["损失函数", "FocalLoss (γ=2)", "FocalLoss (γ=2)", "相同"],
    ["训练结果", "Val F1 97.80%", "Loss=NaN / F1 0.8869(退化)", "8.21 训练失败"],
]
t2 = add_table(doc, cfg, widths=[1.1, 1.6, 1.9, 1.7])
shade_header(t2)

# ================= 三、数据集分析 =================
add_heading(doc, "三、数据集分析", 1)
add_para(doc, "数据集：competition/train.parquet，总计 2,056,871 条 SOC 日志，3 类标签分布极度不平衡。",
         11, False, "justify", CN_FONT, indent=True)
add_para(doc, "表3-1：数据集类别分布与划分", 10.5, True, "center", CN_HEAD)
ds = [
    ["类别", "数量", "占比", "Train (100%)", "Valid (100%)"],
    ["Benign", "1,899,723", "92.4%", "1,899,723", "1,899,723"],
    ["Suspicious", "45,420", "2.2%", "45,420", "45,420"],
    ["Malicious", "111,728", "5.4%", "111,728", "111,728"],
    ["合计", "2,056,871", "100%", "2,056,871", "2,056,871"],
]
t3 = add_table(doc, ds, widths=[1.4, 1.2, 1.0, 1.4, 1.4])
shade_header(t3)
add_figure(doc, os.path.join(CHART_DIR, "fig3_1_class_dist.png"), "图3-1  数据集类别分布", 4.6)
add_para(doc, "关键问题：本次运行 Train 与 Valid 均为 100%（2,056,871 条），二者完全一致。"
              "验证集等同于训练集，无法反映模型在未见数据上的泛化能力；同时该划分覆盖了 8.19 报告所描述的 "
              "85/15 分层设置，属配置回退错误。", 11, False, "justify", CN_FONT, indent=True)

# ================= 四、训练过程 =================
add_heading(doc, "四、训练过程", 1)
add_heading(doc, "4.1 对比学习阶段（Phase 1：NTXent）", 2)
add_para(doc, "对比学习阶段共 10 轮。损失从 4.6733 缓慢降至 4.5769（Ep8），降幅仅 2.1%，"
              "且 Ep9、Ep10 变为 NaN，训练数值发散。", 11, False, "justify", CN_FONT, indent=True)
add_para(doc, "表4-1：对比学习各 Epoch 损失", 10.5, True, "center", CN_HEAD)
cl = [
    ["Epoch", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10"],
    ["Loss", "4.6733", "4.6160", "4.6066", "4.6002", "4.5859", "4.5866", "4.5868", "4.5769", "NaN", "NaN"],
]
t4 = add_table(doc, cl, widths=[0.75] + [0.55]*10)
shade_header(t4)
add_figure(doc, os.path.join(CHART_DIR, "fig4_1_contrastive_loss.png"), "图4-1  对比学习损失曲线（Ep9-10 数值发散为 NaN）", 5.8)
add_para(doc, "观察与风险：", 11, True)
num_para(doc, "(1) 损失前 8 轮下降微小并进入平台期（约 4.58），说明 32D + PCA(16) 仅保留 38.8% 方差，特征表示区分度不足；")
num_para(doc, "(2) Ep9 起损失变为 NaN，说明 NTXent 训练出现数值溢出/不稳定（如 logits 过大导致 exp 溢出），"
              "产生的 NaN 权重将污染后续监督阶段。")

add_heading(doc, "4.2 监督学习阶段（Phase 2：3-class FocalLoss）", 2)
add_para(doc, "监督学习阶段自第 1 epoch 起训练损失即为 NaN，ACC 恒为 0.9236。每轮验证结果完全相同："
              "全部样本被预测为 Benign。第 7 轮触发 Early Stop（6 次无提升），保存的“最佳”模型为 Ep1（同为退化模型）。",
         11, False, "justify", CN_FONT, indent=True)
add_para(doc, "表4-2：监督学习各 Epoch 关键指标（Ep1-7 结果一致）", 10.5, True, "center", CN_HEAD)
sup = [
    ["Epoch", "Train Loss", "Train/Val ACC", "Weighted F1", "Suspicious F1", "Malicious F1", "Status"],
    ["1", "NaN", "0.9236", "0.8869", "0.0000", "0.0000", "★ BEST(退化)"],
    ["2", "NaN", "0.9236", "0.8869", "0.0000", "0.0000", "No improve"],
    ["3", "NaN", "0.9236", "0.8869", "0.0000", "0.0000", "No improve"],
    ["4", "NaN", "0.9236", "0.8869", "0.0000", "0.0000", "No improve"],
    ["5", "NaN", "0.9236", "0.8869", "0.0000", "0.0000", "No improve"],
    ["6", "NaN", "0.9236", "0.8869", "0.0000", "0.0000", "No improve"],
    ["7", "NaN", "0.9236", "0.8869", "0.0000", "0.0000", "Early Stop"],
]
t5 = add_table(doc, sup, widths=[0.6, 0.85, 1.0, 0.95, 1.0, 1.0, 0.9])
shade_header(t5)
add_figure(doc, os.path.join(CHART_DIR, "fig4_2_perclass_f1.png"), "图4-2  监督阶段每类 F1（Suspicious / Malicious 完全失效）", 5.4)
add_para(doc, "关键发现：", 11, True)
num_para(doc, "(1) 训练损失全程 NaN：模型从未获得有效梯度信号，无法学习任何分类边界；")
num_para(doc, "(2) ACC=0.9236 仅为多数类占比：模型把全部样本判为 Benign，该值仅反映类别先验，不代表任何学习能力；")
num_para(doc, "(3) Suspicious / Malicious 完全失效：TPR=0，F1=0；")
num_para(doc, "(4) Early Stop 无法挽救：触发时“最优”Ep1 亦为退化模型，早停只是提前结束了无效训练。")

# ================= 五、评估结果分析 =================
add_heading(doc, "五、评估结果分析", 1)
add_heading(doc, "5.1 指标概览", 2)
add_para(doc, "表5-1：最终评估指标", 10.5, True, "center", CN_HEAD)
ov = [
    ["指标", "数值", "解读"],
    ["ACC", "0.9236", "等于 Benign 占比 92.4%，为多数类占优的假象"],
    ["Macro F1", "0.3201", "3 类平均，其中 2 类为 0，严重偏低"],
    ["Weighted F1", "0.8869", "被 Benign 主导（0.9603），掩盖少数类失效"],
    ["AUC", "N/A", "无法计算（仅输出单类，无正负序区分）"],
]
t6 = add_table(doc, ov, widths=[1.3, 1.1, 4.1])
shade_header(t6)
add_para(doc, "表5-2：每类详细指标", 10.5, True, "center", CN_HEAD)
per = [
    ["类别", "PRE", "SEN", "F1", "TPR", "FPR", "TNR", "FNR"],
    ["Benign", "0.9236", "1.0000", "0.9603", "1.0000", "1.0000", "0.0000", "0.0000"],
    ["Suspicious", "0.0000", "0.0000", "0.0000", "0.0000", "0.0000", "1.0000", "1.0000"],
    ["Malicious", "0.0000", "0.0000", "0.0000", "0.0000", "0.0000", "1.0000", "1.0000"],
    ["Macro", "0.3079", "0.3333", "0.3201", "0.3333", "0.3333", "0.6667", "0.6667"],
    ["Weighted", "0.8530", "0.9236", "0.8869", "0.9236", "0.9236", "0.0764", "0.0764"],
]
t7 = add_table(doc, per, widths=[1.0] + [0.75]*7)
shade_header(t7)
add_para(doc, "分析：Benign 的 FPR=1.0 值得高度警觉——因为所有 Suspicious/Malicious 都被预测为 Benign，"
              "系统对真实威胁的漏报为 100%。Suspicious 45,420 条与 Malicious 111,728 条被全部漏检。",
         11, False, "justify", CN_FONT, indent=True)

add_heading(doc, "5.2 混淆矩阵分析", 2)
add_para(doc, "表5-3：混淆矩阵（Pred vs True）", 10.5, True, "center", CN_HEAD)
cm2 = [
    ["", "Pred_Benign", "Pred_Suspicious", "Pred_Malicious"],
    ["True_Benign", "1,899,723", "0", "0"],
    ["True_Suspicious", "45,420", "0", "0"],
    ["True_Malicious", "111,728", "0", "0"],
]
t8 = add_table(doc, cm2, widths=[1.4, 1.6, 1.6, 1.6])
shade_header(t8)
add_figure(doc, os.path.join(CHART_DIR, "fig5_1_confusion.png"), "图5-1  混淆矩阵热力图（全部预测落在 Benign 列）", 4.6)
add_para(doc, "三个视角解读：", 11, True)
num_para(doc, "(1) 原始数值：只有 Benign 列非零，Pred_Suspicious、Pred_Malicious 全为 0，模型从未输出 Suspicious 或 Malicious；")
num_para(doc, "(2) 行归一化（Recall）：Benign 100%，Suspicious 0%，Malicious 0%；")
num_para(doc, "(3) 列归一化（Precision）：Benign 100%，Suspicious/Malicious 无定义（0 预测）。")
add_para(doc, "结论：这是一种“检测盲区”型失效，对安全系统而言是最危险的结果——所有真实攻击都被当作正常流量放行。",
         11, True, "justify", CN_FONT, indent=True)

# ================= 六、异常与根因诊断 =================
add_heading(doc, "六、异常与根因诊断", 1)
add_para(doc, "针对本次运行，按异常现象、根因、影响与建议逐项诊断：", 11, False, "justify", CN_FONT, indent=True)
issues = [
    ["编号", "异常现象", "根因", "影响", "建议"],
    ["6.1", "数据划分 Train=Valid=100%", "load_data_nosplit() 返回 train 与 valid 均为全量数据",
     "验证集等同训练集，指标虚高且无泛化意义",
     "恢复 85/15 分层，保证验证集 unseen"],
    ["6.2", "PCA 仅保留 38.8% 方差", "TF-IDF vocab=500 压缩到 16D，丢失约 61.2% 信息",
     "特征区分度不足，对比损失停滞于 4.58 平台期",
     "增大 PCA 维度或改用更高信息保留的表示"],
    ["6.3", "对比学习 Ep9-10 损失 NaN", "NTXent logits 过大导致 exp 溢出/数值不稳定",
     "NaN 权重污染后续监督阶段",
     "温度调整、梯度裁剪、log-sum-exp 数值稳定化"],
    ["6.4", "监督学习 Loss=NaN，退化到多数类", "NaN 权重 + FocalLoss 叠加使梯度为 NaN；无有效信号时 argmax 恒为 Benign",
     "Suspicious/Malicious 召回 0%，AUC 不可计算",
     "修复 NaN 源头后重训，并在训练中监控 loss/梯度"],
    ["6.5", "W&B 步数回绕（step bug）", "对比学习将 step 推至 10，监督阶段从 step 1 重新开始",
     "监督指标在 W&B 中不可见（被忽略）",
     "阶段间复位 step，或使用统一递增步数"],
    ["6.6", "自动反频率α[0.36,15.09,6.14]", "代码 bincount 自动反频率加权并 clamp 至 20，覆盖硬编码 [1,20,8]",
     "与 8.19 已废弃的“灾难性失效”配置一致",
     "回退硬编码 α=[1,20,8]，若要自动α需验证收敛性"],
]
t9 = add_table(doc, issues, widths=[0.55, 1.3, 1.7, 1.4, 1.5], font_size=8.5)
shade_header(t9)

# ================= 七、风险与影响 =================
add_heading(doc, "七、风险与影响", 1)
num_para(doc, "(1) 模型不可用：当前 final.pth / encoder.pkl 为退化的多数类模型，不能用于任何 SOC 检测；")
num_para(doc, "(2) 覆盖风险：本次保存目录 models/transcl_v4_819_full 与 8.19 成功运行相近，存在覆盖已知有效权重的风险，建议先备份或重命名；")
num_para(doc, "(3) 安全盲区：Suspicious 与 Malicious 100% 漏报，比高误报更危险——真实攻击被全部放行；")
num_para(doc, "(4) 验证失效：100%/100% 划分使所有验证指标失实，不能作为性能依据。")

# ================= 八、实验结论 =================
add_heading(doc, "八、实验结论", 1)
num_para(doc, "(1) 本次 8.21 运行训练失败：对比阶段 Ep9-10 输出 NaN、监督阶段全程 Loss=NaN，"
               "模型退化为仅预测多数类 Benign 的哑分类器（ACC=0.9236 恰为 Benign 占比，Suspicious/Malicious F1=0，AUC=N/A）。")
num_para(doc, "(2) 根因指向四点：数据划分错误（100%/100% 泄漏）、特征 PCA 信息损失（38.8% 方差）、"
               "对比学习数值发散（NaN）、以及自动反频率 α[0.36,15.09,6.14]（与 8.19 已废弃配置等价）。")
num_para(doc, "(3) Early Stop 无法解决“模型从未学到东西”的问题：其在 Ep7 触发时保存的“最优”Ep1 本身即为退化模型。")
num_para(doc, "(4) 8.19 回退版（85/15 + 硬编码 α[1,20,8]）是当前已知有效配置，8.21 的关键改动全部偏离了该配置。")
num_para(doc, "(5) W&B 步数回绕属日志层缺陷，不影响训练，但导致监督指标不可见，需修复以保留可审计记录。")

# ================= 九、总结与建议 =================
add_heading(doc, "九、总结与建议", 1)
add_para(doc, "一句话结论：8.21 运行因数据泄漏 + NaN 发散 + 退化的 α 配置而失败，其产物不可用于安全检测。", 11, True)
add_para(doc, "关键数值：", 11, True)
bullet(doc, "数据集：2,056,871（Benign 92.4% / Suspicious 2.2% / Malicious 5.4%）")
bullet(doc, "划分：Train = Valid = 100%（数据泄漏）")
bullet(doc, "对比学习：4.6733 → 4.5769（Ep8），Ep9-10 = NaN")
bullet(doc, "监督学习：Loss = NaN，ACC = 0.9236，Weighted F1 = 0.8869")
bullet(doc, "少数类：Suspicious F1 = 0，Malicious F1 = 0，AUC = N/A")
bullet(doc, "最佳模型：Ep1（退化），Early Stop p=6 于 Ep7 触发")
add_para(doc, "下一步建议：", 11, True)
num_para(doc, "(1) 回退到 8.19 已验证配置（85/15 + α[1,20,8]），优先恢复有效基线；")
num_para(doc, "(2) 为 NTXent / 监督训练加入数值稳定与 NaN 防护，训练中持续监控 loss 与梯度；")
num_para(doc, "(3) 增大特征维度（64D/128D）或提升 PCA 保留方差，改善对比学习特征表示；")
num_para(doc, "(4) 使用正确的分层划分，避免 100%/100% 数据泄漏；")
num_para(doc, "(5) 修复 W&B 阶段间步数回绕，保证日志可审计；")
num_para(doc, "(6) 备份/隔离 models/transcl_v4_819_full 目录，避免覆盖已知好模型。")

doc.save(OUT)
print("SAVED:", OUT, os.path.getsize(OUT))
